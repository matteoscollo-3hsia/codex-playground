from __future__ import annotations

from pathlib import Path
from datetime import date
import json
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from markdown_it import MarkdownIt

PLACEHOLDER = "{{CONTENT}}"


def render_primer_docx(
    md_path: str, out_docx_path: str, template_path: str | None = None
) -> None:
    md_file = Path(md_path)
    out_file = Path(out_docx_path)
    template_file = Path(template_path) if template_path else None

    markdown_text = md_file.read_text(encoding="utf-8")
    doc, writer = _init_document(template_file)
    _apply_placeholder_replacements(doc, md_file)
    _apply_style_profile(doc)
    tokens = _markdown_parser().parse(markdown_text)
    _render_tokens(tokens, doc, writer)

    out_file.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_atomic(doc, out_file)


class _DocWriter:
    def __init__(
        self,
        doc: Document,
        insert_after: Paragraph | None,
        reuse_first: Paragraph | None = None,
    ) -> None:
        self.doc = doc
        self.insert_after = insert_after
        self.reuse_first = reuse_first

    def add_paragraph(self, style: str | None = None) -> Paragraph:
        if self.reuse_first is not None:
            paragraph = self.reuse_first
            self.reuse_first = None
        elif self.insert_after is None:
            paragraph = self.doc.add_paragraph()
        else:
            paragraph = _insert_paragraph_after(self.insert_after)
        if style:
            try:
                paragraph.style = style
            except KeyError:
                pass
        self.insert_after = paragraph
        return paragraph

    def add_table(self, rows: int, cols: int):
        table = self.doc.add_table(rows=rows, cols=cols)
        if self.insert_after is not None:
            self.insert_after._p.addnext(table._tbl)
        new_p = OxmlElement("w:p")
        table._tbl.addnext(new_p)
        paragraph_after = Paragraph(new_p, table._parent)
        self.reuse_first = paragraph_after
        self.insert_after = paragraph_after
        return table


def _markdown_parser() -> MarkdownIt:
    md = MarkdownIt("commonmark", {"breaks": False, "html": False})
    md.enable("table")
    return md


def _init_document(template_path: Path | None) -> tuple[Document, _DocWriter]:
    if template_path:
        _ensure_template(template_path)
        if template_path.exists() and template_path.is_dir():
            raise ValueError(f"Template path is a directory: {template_path}")
        doc = Document(str(template_path))
        placeholder_paragraph = _find_placeholder_paragraph(doc, PLACEHOLDER)
        if placeholder_paragraph is not None:
            _remove_placeholder_text(placeholder_paragraph, PLACEHOLDER)
            reuse_first = None
            if not placeholder_paragraph.text.strip():
                reuse_first = placeholder_paragraph
            return doc, _DocWriter(
                doc, insert_after=placeholder_paragraph, reuse_first=reuse_first
            )
        return doc, _DocWriter(doc, insert_after=None)
    doc = Document()
    return doc, _DocWriter(doc, insert_after=None)


def _ensure_template(template_path: Path) -> None:
    if template_path.exists():
        return
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("{{COMPANY}} – Commercial Primer")
    doc.add_paragraph("Date: {{DATE}}")
    doc.add_page_break()
    doc.add_paragraph(PLACEHOLDER)
    _save_docx_atomic(doc, template_path)


def _find_placeholder_paragraph(doc: Document, placeholder: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            return paragraph
    return None


def _remove_placeholder_text(paragraph: Paragraph, placeholder: str) -> None:
    if placeholder not in paragraph.text:
        return
    paragraph.text = paragraph.text.replace(placeholder, "")


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _save_docx_atomic(doc: Document, path: Path) -> None:
    tmp_path = path.with_suffix(path.suffix + ".tmp")
    doc.save(str(tmp_path))
    tmp_path.replace(path)


def _render_tokens(tokens: list[Any], doc: Document, writer: _DocWriter) -> None:
    list_stack: list[str] = []
    in_list_item = 0
    i = 0
    while i < len(tokens):
        token = tokens[i]
        if token.type == "bullet_list_open":
            list_stack.append("bullet")
            i += 1
            continue
        if token.type == "ordered_list_open":
            list_stack.append("ordered")
            i += 1
            continue
        if token.type in ("bullet_list_close", "ordered_list_close"):
            if list_stack:
                list_stack.pop()
            i += 1
            continue
        if token.type == "list_item_open":
            in_list_item += 1
            i += 1
            continue
        if token.type == "list_item_close":
            in_list_item = max(0, in_list_item - 1)
            i += 1
            continue

        if token.type == "heading_open":
            level = int(token.tag[1]) if token.tag and token.tag.startswith("h") else 1
            inline = tokens[i + 1] if i + 1 < len(tokens) and tokens[i + 1].type == "inline" else None
            runs = _inline_runs(inline)
            style = _heading_style(doc, level)
            _add_paragraph(writer, style, runs)
            i = _skip_to(tokens, i, "heading_close") + 1
            continue

        if token.type == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < len(tokens) and tokens[i + 1].type == "inline" else None
            runs = _inline_runs(inline)
            style = _paragraph_style(doc, list_stack, in_list_item)
            _add_paragraph(
                writer,
                style,
                runs,
                list_indent=in_list_item > 0 and list_stack,
            )
            i = _skip_to(tokens, i, "paragraph_close") + 1
            continue

        if token.type == "hr":
            _add_paragraph(writer, _paragraph_style(doc, [], 0), [])
            i += 1
            continue

        if token.type == "fence":
            _add_code_block(writer, token.content, _paragraph_style(doc, [], 0))
            i += 1
            continue

        if token.type == "table_open":
            table_rows, end_index = _parse_table(tokens, i)
            if table_rows:
                table = _render_table(writer, doc, table_rows)
                _apply_table_profile(table, doc)
            i = end_index + 1
            continue

        i += 1


def _skip_to(tokens: list[Any], start_index: int, end_type: str) -> int:
    i = start_index + 1
    while i < len(tokens) and tokens[i].type != end_type:
        i += 1
    return i


def _heading_style(doc: Document, level: int) -> str | None:
    if level <= 1:
        return "Title" if _style_exists(doc, "Title") else "Heading 1"
    if level == 2:
        return "Heading 1"
    return "Heading 2"


def _paragraph_style(doc: Document, list_stack: list[str], in_list_item: int) -> str | None:
    if in_list_item > 0 and list_stack:
        list_kind = list_stack[-1]
        if list_kind == "bullet":
            style = "List Bullet"
        else:
            style = "List Number"
        return style
    return "Normal"


def _style_exists(doc: Document, name: str) -> bool:
    try:
        doc.styles[name]
    except KeyError:
        return False
    return True


def _add_paragraph(
    writer: _DocWriter,
    style: str | None,
    runs: list[dict[str, Any]],
    *,
    list_indent: bool = False,
) -> None:
    paragraph = writer.add_paragraph(style)
    if list_indent:
        _apply_list_indent(paragraph)
    _add_runs(paragraph, runs)


def _add_runs(paragraph: Paragraph, runs: list[dict[str, Any]]) -> None:
    for run_spec in runs:
        if run_spec.get("break"):
            paragraph.add_run().add_break()
            continue
        text = run_spec.get("text", "")
        if not text:
            continue
        run = paragraph.add_run(text)
        if run_spec.get("bold"):
            run.bold = True
        if run_spec.get("italic"):
            run.italic = True
        if run_spec.get("code"):
            run.font.name = "Consolas"


def _add_code_block(writer: _DocWriter, content: str, style: str | None) -> None:
    lines = content.rstrip("\n").splitlines()
    if not lines:
        lines = [""]
    for line in lines:
        paragraph = writer.add_paragraph(style)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"


def _inline_runs(inline_token: Any | None) -> list[dict[str, Any]]:
    if inline_token is None or not getattr(inline_token, "children", None):
        return []
    runs: list[dict[str, Any]] = []
    bold = False
    italic = False
    for child in inline_token.children:
        token_type = child.type
        if token_type == "strong_open":
            bold = True
            continue
        if token_type == "strong_close":
            bold = False
            continue
        if token_type == "em_open":
            italic = True
            continue
        if token_type == "em_close":
            italic = False
            continue
        if token_type == "code_inline":
            _append_run(runs, child.content, False, False, True)
            continue
        if token_type == "text":
            _append_run(runs, child.content, bold, italic, False)
            continue
        if token_type == "softbreak":
            _append_run(runs, " ", bold, italic, False)
            continue
        if token_type == "hardbreak":
            runs.append({"break": True})
            continue
        if token_type == "image":
            alt_text = child.content or ""
            placeholder = "TODO: image omitted"
            if alt_text:
                placeholder = f"TODO: image omitted ({alt_text})"
            _append_run(runs, placeholder, bold, italic, False)
            continue
        if token_type in ("link_open", "link_close"):
            continue
    return runs


def _append_run(
    runs: list[dict[str, Any]], text: str, bold: bool, italic: bool, code: bool
) -> None:
    if not text:
        return
    if runs:
        last = runs[-1]
        if (
            not last.get("break")
            and last.get("bold") == bold
            and last.get("italic") == italic
            and last.get("code") == code
        ):
            last["text"] = f"{last.get('text', '')}{text}"
            return
    runs.append({"text": text, "bold": bold, "italic": italic, "code": code})


def _apply_placeholder_replacements(doc: Document, md_file: Path) -> None:
    company_name = _resolve_company_name(md_file)
    replacements = {
        "{{COMPANY}}": company_name or "Unknown Company",
        "{{DATE}}": date.today().isoformat(),
    }

    for paragraph in list(doc.paragraphs):
        if "{{CONTACT}}" in paragraph.text:
            _remove_paragraph(paragraph)
            continue
        _replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    if "{{CONTACT}}" in paragraph.text:
                        _remove_paragraph(paragraph)
                        continue
                    _replace_text_in_paragraph(paragraph, replacements)


def _replace_text_in_paragraph(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    text = paragraph.text
    if not text:
        return
    updated = text
    for key, value in replacements.items():
        if key in updated:
            updated = updated.replace(key, value)
    if updated != text:
        paragraph.text = updated


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._p
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _resolve_company_name(md_file: Path) -> str | None:
    candidate_paths: list[Path] = []
    for parent in [md_file.parent, *md_file.parents]:
        candidate_paths.append(parent / "lead_input.json")
        candidate_paths.append(parent / "_dossier" / "lead_input.json")

    seen: set[Path] = set()
    for path in candidate_paths:
        if path in seen:
            continue
        seen.add(path)
        if not path.exists():
            continue
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        if not isinstance(payload, dict):
            continue
        for key in ("company_name", "client", "company"):
            value = payload.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
    return None


def _apply_style_profile(doc: Document) -> None:
    profiles = {
        "Title": {
            "font_size_pt": 18.0,
            "bold": True,
            "spacing_before_pt": 24.0,
            "spacing_after_pt": 12.0,
            "keep_with_next": True,
        },
        "Heading 1": {
            "font_size_pt": 16.0,
            "bold": True,
            "spacing_before_pt": 24.0,
            "spacing_after_pt": 0.0,
            "keep_with_next": True,
        },
        "Heading 2": {
            "font_size_pt": 14.0,
            "bold": True,
            "spacing_before_pt": 10.0,
            "spacing_after_pt": 0.0,
            "keep_with_next": True,
        },
        "Heading 3": {
            "bold": True,
            "spacing_before_pt": 10.0,
            "spacing_after_pt": 0.0,
            "keep_with_next": True,
        },
        "Normal": {
            "spacing_before_pt": 9.0,
            "spacing_after_pt": 9.0,
        },
    }

    for style_name, profile in profiles.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        font = style.font
        if profile.get("font_size_pt") is not None and font.size is None:
            font.size = Pt(profile["font_size_pt"])
        if profile.get("bold") is not None and font.bold is None:
            font.bold = profile["bold"]
        fmt = style.paragraph_format
        if profile.get("spacing_before_pt") is not None and fmt.space_before is None:
            fmt.space_before = Pt(profile["spacing_before_pt"])
        if profile.get("spacing_after_pt") is not None and fmt.space_after is None:
            fmt.space_after = Pt(profile["spacing_after_pt"])
        if profile.get("keep_with_next") is not None and fmt.keep_with_next is None:
            fmt.keep_with_next = profile["keep_with_next"]


def _apply_list_indent(paragraph: Paragraph) -> None:
    fmt = paragraph.paragraph_format
    if fmt.left_indent is None:
        fmt.left_indent = Pt(18.0)
    if fmt.first_line_indent is None:
        fmt.first_line_indent = Pt(-9.0)


def _apply_table_profile(table, doc: Document) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    if table.columns and len(table.columns) == 2:
        section = doc.sections[0]
        total_width = section.page_width - section.left_margin - section.right_margin
        half_width = int(total_width / 2)
        for col in table.columns:
            col.width = half_width

    if table.rows:
        _set_header_repeat(table.rows[0])


def _set_header_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def _parse_table(tokens: list[Any], start_index: int) -> tuple[list[list[list[dict[str, Any]]]], int]:
    rows: list[list[list[dict[str, Any]]]] = []
    row_cells: list[list[dict[str, Any]]] = []
    i = start_index + 1
    while i < len(tokens):
        token = tokens[i]
        if token.type == "tr_open":
            row_cells = []
        elif token.type in ("th_open", "td_open"):
            runs: list[dict[str, Any]] = []
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                runs = _inline_runs(tokens[i + 1])
            row_cells.append(runs)
        elif token.type == "tr_close":
            rows.append(row_cells)
        elif token.type == "table_close":
            return rows, i
        i += 1
    return rows, i


def _render_table(
    writer: _DocWriter, doc: Document, rows: list[list[list[dict[str, Any]]]]
):
    if not rows:
        return None
    cols = max(len(row) for row in rows)
    table = writer.add_table(rows=len(rows), cols=cols)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            runs = row[c_idx] if c_idx < len(row) else []
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            _add_runs(paragraph, runs)
    return table
